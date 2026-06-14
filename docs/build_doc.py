"""
build_doc.py — Generate Vera Level FX System Documentation Word file.
Run: python docs/build_doc.py
Output: docs/Vera_Level_FX_System_Documentation.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), 'Vera_Level_FX_System_Documentation.docx')

# ── Colour palette ──────────────────────────────────────────────
NAVY   = RGBColor(0x01, 0x0E, 0x1F)
GOLD   = RGBColor(0xF0, 0xC0, 0x40)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
MUTED  = RGBColor(0x44, 0x44, 0x66)
GREEN  = RGBColor(0x00, 0x80, 0x50)
RED    = RGBColor(0xCC, 0x33, 0x33)

doc = Document()

# ── Page margins ────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)


# ── Helper functions ────────────────────────────────────────────

def set_run_font(run, name='Calibri', size=11, bold=False, italic=False,
                 color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_run_font(run, name='Calibri', size=18, bold=True, color=NAVY)
    # Gold underline bar via bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'F0C040')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=13, bold=True, color=NAVY)
    return p


def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=11, bold=True, color=MUTED)
    return p


def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def code_block(lines):
    """Add a shaded monospace block."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        # Light grey shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        pPr.append(shd)
        run = p.add_run(line if line else ' ')
        set_run_font(run, name='Courier New', size=9, color=DARK)
    doc.add_paragraph()   # trailing space


def table_2col(rows, col_widths=(6, 10), header=None):
    """Simple 2-column table."""
    t = doc.add_table(rows=len(rows) + (1 if header else 0), cols=2)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    col_w = [Cm(w) for w in col_widths]
    for i, row in enumerate(t.rows):
        for j, cell in enumerate(row.cells):
            cell.width = col_w[j]

    offset = 0
    if header:
        for j, h in enumerate(header):
            cell = t.rows[0].cells[j]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(h)
            set_run_font(run, size=10, bold=True, color=WHITE)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '010E1F')
            tcPr.append(shd)
        offset = 1

    for i, (a, b) in enumerate(rows):
        row = t.rows[i + offset]
        run_a = row.cells[0].paragraphs[0].add_run(a)
        run_b = row.cells[1].paragraphs[0].add_run(b)
        set_run_font(run_a, size=10, bold=True)
        set_run_font(run_b, size=10)
        # Alternate row shading
        if (i + offset) % 2 == 1:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F7F7F7')
                tcPr.append(shd)

    doc.add_paragraph()
    return t


def note_box(text, color_hex='FFF3CD'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)
    run = p.add_run(text)
    set_run_font(run, size=10, italic=True)
    return p


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('VERA LEVEL FX')
set_run_font(run, name='Calibri', size=32, bold=True, color=NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('System Documentation & End-to-End Posting Flow')
set_run_font(run, size=16, color=MUTED)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Algorithmic Forex Trading  ·  IC Markets  ·  Myfxbook Verified')
set_run_font(run, size=11, italic=True, color=MUTED)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('@veralevel.fx  ·  Singapore')
set_run_font(run, size=11, color=MUTED)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Prepared by Pandian  ·  June 2026')
set_run_font(run, size=10, color=MUTED)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 1. SYSTEM OVERVIEW
# ════════════════════════════════════════════════════════════════

heading1('1. System Overview')

body(
    'Vera Level FX is a fully automated algorithmic forex trading system running on a live IC Markets '
    'account. The system uses a quantitative multi-pair grid strategy across EURUSD, GBPJPY, XAUUSD, '
    'and USDJPY. All trades are publicly verified on Myfxbook in real time.'
)
body(
    'This document covers two interconnected systems:'
)
bullet('myfxbook-mcp — the local data pipeline that scrapes live account data daily and writes it to disk.')
bullet('vera-level-fx — the Instagram automation pipeline that reads that data and posts graphics to @veralevel.fx.')

heading2('1.1  High-Level Architecture')

body('The full system spans three machines/services:')

table_2col(
    [
        ('Windows Laptop (home IP)',
         'Runs daily Myfxbook scraper at 08:00 SGT. Bypasses Cloudflare via residential IP. '
         'Pushes fresh data to both GitHub repos.'),
        ('GitHub (vera-level-fx repo)',
         'Stores the data snapshot (data/vera-snapshot.json) and all Instagram code. '
         'GitHub Actions generates and posts images on schedule.'),
        ('Instagram / Meta API',
         'Receives the final 1080×1080 PNG and caption via Meta Graph API. '
         'Post appears on @veralevel.fx automatically.'),
    ],
    col_widths=(5, 11),
    header=('Component', 'Role'),
)

heading2('1.2  Repository Structure')

table_2col(
    [
        ('myfxbook-mcp/', 'Local scraper project (private, Windows laptop only)'),
        ('myfxbook-mcp/src/runner.js', 'Daily entry point — scrapes both accounts, sends Telegram, pushes data'),
        ('myfxbook-mcp/src/scraper.js', 'Puppeteer + stealth Chrome scraper for Myfxbook public page'),
        ('myfxbook-mcp/src/fetcher.js', 'Myfxbook REST API client (authenticated — open trades, history, daily gain)'),
        ('myfxbook-mcp/public/vera-snapshot.json', 'Last scraped Vera Level data (committed to GitHub Pages repo)'),
        ('vera-level-fx/data/vera-snapshot.json', 'Mirror of above — what the Instagram pipeline reads'),
        ('vera-level-fx/instagram/run.py', 'Instagram pipeline entry point — decides post type, generates image, publishes'),
        ('vera-level-fx/instagram/generate_status.py', 'Daily live-positions card renderer (notepad background, ink overlay)'),
        ('vera-level-fx/instagram/assets/qr-icmarkets.jpg', 'IC Markets referral QR code — stamped onto every post image before publishing'),
        ('vera-level-fx/instagram/generate.py', 'Weekly performance, monthly P&L chart, and win-rate card renderers'),
        ('vera-level-fx/instagram/generate_edu.py', 'Educational post renderers (Risk, Pair Spotlight, Trade Setup)'),
        ('vera-level-fx/instagram/edu_content.py', 'Educational content library — 5 risk rules, 4 pairs, 4 setups'),
        ('vera-level-fx/instagram/captions.py', 'Instagram caption templates for all post types'),
        ('vera-level-fx/instagram/post.py', 'Meta Graph API publisher'),
        ('vera-level-fx/data/edu-counter.json', '6-week educational post rotation counter'),
        ('vera-level-fx/.github/workflows/insta-post.yml', 'GitHub Actions workflow — cron schedule + workflow_dispatch trigger'),
    ],
    col_widths=(7, 9),
    header=('File / Directory', 'Purpose'),
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 2. DATA PIPELINE — MYFXBOOK SCRAPER
# ════════════════════════════════════════════════════════════════

heading1('2. Data Pipeline — Myfxbook Scraper (myfxbook-mcp)')

heading2('2.1  What It Does')

body(
    'Every weekday at 08:00 SGT, a Windows Task Scheduler job runs '
    '"node src/runner.js --push" from the myfxbook-mcp project directory. '
    'This single command:'
)
bullet('Launches a stealth headless Chrome session via Puppeteer.')
bullet('Scrapes the TOL LANGIT ETF GOLD public Myfxbook page (no login required).')
bullet('Logs in to Myfxbook and calls the REST API to fetch Vera Level account stats, open trades, trade history, and daily gain data.')
bullet('Merges authenticated DOM-scraped lifetime stats (win rate, trade count, profit factor) with the REST API data.')
bullet('Sends a Telegram report to the operator chat.')
bullet('Writes vera-snapshot.json to both the myfxbook-mcp GitHub Pages repo and the vera-level-fx repo.')
bullet('Commits and pushes both repos — triggering a Vercel site redeploy and making fresh data available to GitHub Actions.')

heading2('2.2  Why the Scraper Runs Locally (Not in the Cloud)')

body(
    'Myfxbook uses Cloudflare bot protection. Cloud datacenter IPs (GitHub Actions, AWS, GCP) are '
    'blocked outright. The scraper must run from a residential ISP connection. The Windows laptop at '
    'home passes Cloudflare because:'
)
bullet('The IP is a real residential address, not a known datacenter range.')
bullet('puppeteer-extra-plugin-stealth spoofs browser fingerprints (user agent, canvas, WebGL, etc.).')
bullet('The real Chrome binary is used, not Chromium, which avoids additional detection vectors.')

note_box(
    '⚠  If you run the scraper with a VPN active, Cloudflare will block it. Always disable VPN '
    'before the 08:00 SGT scheduled run.',
    'FFF3CD'
)

heading2('2.3  Data Captured')

table_2col(
    [
        ('balance', 'Current account balance in USD'),
        ('equity', 'Current equity (balance + floating P&L) in USD'),
        ('daily', 'Daily gain percentage (from REST API)'),
        ('monthly', '30-day rolling average gain percentage'),
        ('gain', 'Total gain % since account inception'),
        ('drawdown', 'Maximum balance drawdown percentage'),
        ('profit', 'Total realised profit in USD'),
        ('pips', 'Total pips accumulated across all closed trades'),
        ('profitFactor', 'Gross profit / gross loss ratio'),
        ('trades', 'Total number of closed trades'),
        ('winRate', 'Percentage of trades closed in profit'),
        ('avgWin / avgLoss', 'Average win and loss in USD'),
        ('openTrades[]', 'Array of all currently open positions (symbol, direction, profit, pips, open time)'),
        ('openOrders[]', 'Pending orders (buy/sell stops and limits)'),
        ('history[]', 'Up to 40 most recent closed trades with full P&L detail'),
        ('dailyGain[]', 'Cumulative daily gain time series from account start'),
    ],
    col_widths=(5, 11),
    header=('Field', 'Description'),
)

heading2('2.4  Data Format — vera-snapshot.json')

body('The snapshot file written to vera-level-fx/data/vera-snapshot.json follows this structure:')

code_block([
    '{',
    '  "account": {',
    '    "balance": 2978.24,',
    '    "equity":  3054.55,',
    '    "daily":  -0.06,',
    '    "monthly": -1.67,',
    '    "gain":    -6.10,',
    '    "drawdown": 51.04,',
    '    "profitFactor": 0.94,',
    '    "pips":    31819.8,',
    '    "trades":  673,',
    '    "winRate": 76',
    '  },',
    '  "openTrades": [',
    '    {',
    '      "openTime": "06/11/2026 05:15",   // MM/DD/YYYY HH:MM',
    '      "symbol":   "XAUUSD",',
    '      "action":   "Buy",',
    '      "profit":   136.17,               // float USD',
    '      "pips":     13617.0               // float (XAU = 0.01/pip)',
    '    }, ...',
    '  ],',
    '  "dailyGain": [["02/20/2026", 0, 0], ["02/24/2026", 2.8, 27.96], ...],',
    '  "fetchedAt": "2026-06-13T10:19:15.881Z"',
    '}',
])

heading2('2.5  Windows Task Scheduler Configuration')

table_2col(
    [
        ('Task name',    'TOL LANGIT Daily Report'),
        ('Schedule',     'Mon–Fri at 08:00 SGT (UTC+8)'),
        ('Command',      'node "C:\\Users\\USER\\Downloads\\myfxbook-mcp\\src\\runner.js" --push'),
        ('Working dir',  'C:\\Users\\USER\\Downloads\\myfxbook-mcp'),
        ('On failure',   'Retry once after 5 minutes'),
        ('Missed run',   'StartWhenAvailable — runs when laptop next turns on'),
        ('Timeout',      '10 minutes'),
    ],
    col_widths=(4, 12),
    header=('Setting', 'Value'),
)

body('Manual trigger command (PowerShell):')
code_block(['Start-ScheduledTask -TaskName "TOL LANGIT Daily Report"'])

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 3. INSTAGRAM AUTOMATION PIPELINE
# ════════════════════════════════════════════════════════════════

heading1('3. Instagram Automation Pipeline (vera-level-fx)')

heading2('3.1  Post Schedule')

body(
    'GitHub Actions triggers the Instagram pipeline on a fixed weekly schedule. '
    'All times are 09:00 SGT (01:00 UTC).'
)

table_2col(
    [
        ('Monday',        'Weekly Performance card — balance, total gain, win rate, profit factor, drawdown'),
        ('Tuesday',       'Educational post (rotating 6-week library) — Risk Management rules'),
        ('Wednesday',     'Daily Live Positions card — open trades on notepad photo background'),
        ('Thursday',      'Educational post (rotating 6-week library) — Pair Spotlight or Trade Setup'),
        ('Friday',        'Daily Live Positions card — open trades update'),
        ('1st of month',  'Monthly P&L bar chart — 12-month rolling breakdown'),
    ],
    col_widths=(4, 12),
    header=('Day', 'Post Type'),
)

note_box(
    '📌  The "daily" post type is also used for Wednesday and Friday. '
    'Wednesday posts show mid-week open position snapshot; Friday posts show week-close positions.',
    'E8F4FD'
)

heading2('3.2  Post Type Decision Logic')

body('run.py determines the post type via the following priority order:')
bullet('If today is the 1st or 2nd of the month → monthly P&L chart.')
bullet('If Monday → weekly performance card.')
bullet('If Tuesday or Thursday → educational post (6-week rotating library).')
bullet('If Wednesday or Friday → daily live positions card.')
bullet('POST_TYPE environment variable overrides all of the above (used for manual triggers).')

heading2('3.3  Image Generation')

body(
    'All images are 1080×1080 pixels at 100 DPI, rendered with Python\'s Matplotlib library. '
    'No external design tools or paid services are used — every post is generated entirely in code.'
)

heading3('Daily Live Positions Card (generate_status.py)')

body(
    'Background: a photograph of an open notebook on a desk with a MacBook showing a trading chart. '
    'The table spans the full image width (x 0.05–0.95), centred at the midpoint.'
)
bullet('A dark semi-transparent overlay (alpha 0.62) covers the laptop/keyboard area at the top of the image.')
bullet('Branding strip: gold bar with "VERA LEVEL FX", live date/time, equity/balance/daily metrics.')
bullet('Trade data is rendered in INK colour (#0A1628) directly onto the white notebook page.')
bullet('Up to 5 open positions shown: PAIR · DIR · P&L (USD) · PIPS · ENTRY (open price).')
bullet('ENTRY price formatted as 4 decimal places for FX pairs (e.g. 1.1829), 2 decimal places for instruments ≥ 10 (e.g. 4089.47 for XAUUSD).')
bullet('Large pip values (>999) are abbreviated: e.g. 13,617 pips → "+13.6k".')
bullet('Stats footer: win rate, profit factor, total pips — centred across full width.')
bullet('Duplicate positions (same symbol + time + action) are deduplicated before display.')

heading3('Weekly Performance Card (generate.py)')

body(
    'Dark navy background with 6 metric cards in a 2×3 grid: Account Balance, Total Gain, '
    'Monthly Average, Win Rate, Profit Factor, and Drawdown. '
    'Green for positive values, red for negative, gold for balance.'
)

heading3('Monthly P&L Chart (generate.py)')

body(
    'A 12-month rolling bar chart of daily gain data aggregated by month. '
    'Green bars for positive months, red for negative. Values labelled on each bar.'
)

heading3('Win Rate / Trust Card (generate.py)')

body(
    'A gold donut chart showing win rate percentage, with three stat cards below '
    '(Profit Factor, Total Gain, Trades). Designed to build social proof and trust.'
)

heading3('IC Markets Referral QR Code Overlay (run.py — overlay_qr)')

body(
    'After every image is generated (or retrieved from the buffer), run.py stamps the IC Markets '
    'referral QR code into the bottom-right corner of the PNG before committing and publishing.'
)
bullet('QR code asset: instagram/assets/qr-icmarkets.jpg (100×100 RGBA source).')
bullet('Resized to ~154px on a 1080 canvas (1/7th of image width) — large enough to scan from a phone screen.')
bullet('Placed on a white semi-transparent background (alpha 235) with 10px padding for contrast on any card colour.')
bullet('Positioned 18px from the right edge and 28px from the bottom — avoids cutting off by Instagram\'s rounded corners.')
bullet('Links to: https://icmarkets.com/?camp=91936 — IB referral campaign #91936.')
bullet('If the QR asset is missing or Pillow fails, the overlay step is skipped silently — the post still publishes.')

heading3('Educational Posts (generate_edu.py)')

body('Three sub-types on a 6-week rotation (posted every Tuesday and Thursday):')
bullet('Risk Management — one of 5 rules (lot sizing, grid spacing, drawdown psychology, diversification, profit factor).')
bullet('Pair Spotlight — deep-dive on EURUSD, GBPJPY, XAUUSD, or USDJPY (session, spread, volatility, edge).')
bullet('Trade Setup Breakdown — step-by-step entry methodology for a specific pair and direction.')

heading2('3.4  Caption Generation (captions.py)')

body(
    'Every image is paired with a caption generated from the same live data. '
    'Captions include the key metrics, a shared call-to-action block, and a fixed set of hashtags.'
)

table_2col(
    [
        ('weekly(account)',               'Balance, gain, monthly avg, win rate, profit factor, pips + CTA + hashtags'),
        ('monthly(account, pnl_dict)',    'Last 6 months breakdown with green/red emoji per month + CTA + hashtags'),
        ('daily_status(account, trades)', 'Equity, balance, daily %, open positions list with profit emoji + CTA + hashtags'),
        ('trust(account)',                'Win rate, profit factor, total gain, pips + credibility copy + CTA + hashtags'),
        ('edu(type, content)',            'Type-specific educational copy + CTA + edu hashtags'),
    ],
    col_widths=(5.5, 10.5),
    header=('Function', 'Content'),
)

heading3('Call-to-Action Block (_CTA)')

body('All five caption functions share a single _CTA constant appended at the end of every caption:')

code_block([
    '📲 Live signals → https://t.me/pandiangk',
    '🌐 Live account → https://vera-level-forex.vercel.app',
    '🏦 Open IC Markets account →',
    'https://icmarkets.com/global/en/?camp=91936',
])

note_box(
    '📌  All three URLs use the full https:// prefix so Instagram renders them as tappable links '
    'on mobile. The IC Markets URL includes the IB referral parameter camp=91936.',
    'E8F4FD'
)

heading2('3.5  Publishing to Instagram (post.py)')

body('The Meta Graph API is used for all Instagram publishing. The flow is:')
bullet('Step 1 — Image is committed to the GitHub repo as a PNG file.')
bullet('Step 2 — The raw GitHub URL is used as the image source (Meta API requires a public URL, not a local file).')
bullet('Step 3 — A 20-second wait allows the GitHub CDN to propagate the file globally.')
bullet('Step 4 — POST to /media endpoint creates a media container with the image URL and caption.')
bullet('Step 5 — POST to /media_publish endpoint publishes the container to @veralevel.fx.')

body('Environment variables required:')
code_block([
    'IG_USER_ID        — Instagram Business account numeric ID',
    'META_ACCESS_TOKEN — Long-lived Page Access Token (never expires if refreshed monthly)',
])

note_box(
    '🔒  Both tokens are stored in GitHub Actions Secrets and are never committed to the repository. '
    'The GitHub Actions workflow injects them as environment variables at runtime.',
    'E8F4FD'
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 4. CONTENT BUFFER
# ════════════════════════════════════════════════════════════════

heading1('4. Content Buffer — Pre-Made Image Queue')

heading2('4.1  What It Is')

body(
    'The content buffer is a set of folders on your laptop where you can drop pre-designed '
    '1080×1080 PNG images during your free time. When the Instagram pipeline runs, it checks '
    'the buffer first. If a buffered image exists for that post type, it uses that instead of '
    'auto-generating one. If the buffer is empty, it falls back to auto-generation as normal.'
)
body(
    'This lets you batch-create high-quality graphics (in Canva, on your phone, or any design '
    'tool) whenever you have spare time, then let the pipeline drain them automatically on schedule.'
)

heading2('4.2  Folder Structure')

body('Location: vera-level-fx/instagram/buffer/')

table_2col(
    [
        ('buffer/daily/',       'Pre-made live position cards — used on Wednesday and Friday posts'),
        ('buffer/weekly/',      'Pre-made weekly performance cards — used on Monday posts'),
        ('buffer/monthly/',     'Pre-made monthly P&L charts — used on 1st of month posts'),
        ('buffer/trust/',       'Pre-made win rate / trust cards — used when post type is "trust"'),
        ('buffer/edu/risk/',    'Pre-made risk management educational posts — Tue/Thu rotation'),
        ('buffer/edu/pairs/',   'Pre-made pair spotlight educational posts — Tue/Thu rotation'),
        ('buffer/edu/setup/',   'Pre-made trade setup breakdown posts — Tue/Thu rotation'),
    ],
    col_widths=(5, 11),
    header=('Folder', 'Purpose'),
)

heading2('4.3  How It Works')

bullet('Drop one or more 1080×1080 PNG files into the matching subfolder.')
bullet('Commit and push to GitHub.')
bullet('On the next scheduled post for that type, run.py calls pop_buffer() before generating.')
bullet('pop_buffer() picks the alphabetically oldest PNG from the folder (FIFO).')
bullet('The file is immediately moved to instagram/posts/YYYY-MM-DD-{type}-buffered.png — it will never be used twice.')
bullet('Caption is always generated live from the current vera-snapshot.json data — only the image is buffered.')
bullet('If the buffer folder is empty, the pipeline generates the image automatically as usual.')

note_box(
    '💡  To control the order in which buffered images are used, prefix filenames with a number: '
    '"01_gold-setup.png", "02_eurusd-chart.png". They will be consumed in alphabetical order.',
    'E8F4FD'
)

heading2('4.4  Step-by-Step: Adding an Image to the Buffer')

body('From your laptop:')
code_block([
    '# 1. Copy your PNG into the correct subfolder:',
    '#    vera-level-fx\\instagram\\buffer\\daily\\my-chart-June.png',
    '',
    '# 2. Stage and commit:',
    'cd "C:\\Users\\USER\\Downloads\\04 - FX Trading\\vera-level-fx"',
    'git add instagram/buffer/',
    'git commit -m "buffer: add daily chart for June"',
    'git push',
    '',
    '# 3. Done. The next Wednesday or Friday post will use your image.',
])

heading2('4.5  Checking Buffer Status')

body('See what images are currently queued:')
code_block([
    'ls vera-level-fx\\instagram\\buffer\\daily\\',
    'ls vera-level-fx\\instagram\\buffer\\weekly\\',
    'ls vera-level-fx\\instagram\\buffer\\edu\\risk\\',
    '# etc.',
])

heading2('4.6  Buffer Logic in run.py')

body('The pop_buffer() function in instagram/run.py handles buffer consumption:')
code_block([
    'def pop_buffer(post_type, edu_type=""):',
    '    folder = BUFFER_DIR / post_type          # e.g. buffer/daily/',
    '    if edu_type:',
    '        folder = BUFFER_DIR / "edu" / edu_type  # e.g. buffer/edu/risk/',
    '    pngs = sorted(folder.glob("*.png"))      # alphabetical order',
    '    if not pngs:',
    '        return None                           # buffer empty → auto-generate',
    '    src = pngs[0]                             # oldest file',
    '    dst = OUT_DIR / f"{today}-{post_type}-buffered.png"',
    '    src.rename(dst)                           # move out of buffer',
    '    return dst                                # pipeline uses this path',
])

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 5. EDUCATIONAL CONTENT ROTATION
# ════════════════════════════════════════════════════════════════

heading1('5. Educational Content Rotation (6-Week Cycle)')

heading2('5.1  Rotation Structure')

body(
    'Educational posts are published every Tuesday and Thursday at 09:00 SGT. '
    'A persistent counter in data/edu-counter.json tracks the current position in a fixed '
    '12-slot, 6-week sequence. The counter is committed back to the repo after each post, '
    'so the next post always picks up where the last left off — even if a post fails or is skipped.'
)

table_2col(
    [
        ('Week 1 Tue',  'Risk Rule #01 — Lot Sizing By Balance'),
        ('Week 1 Thu',  'Pair Spotlight — EURUSD'),
        ('Week 2 Tue',  'Trade Setup — EURUSD Long (Liquidity Sweep + Reentry)'),
        ('Week 2 Thu',  'Risk Rule #02 — Grid Spacing Discipline'),
        ('Week 3 Tue',  'Pair Spotlight — XAUUSD'),
        ('Week 3 Thu',  'Trade Setup — XAUUSD Short (Institutional Rejection)'),
        ('Week 4 Tue',  'Risk Rule #03 — Balance vs Equity Drawdown'),
        ('Week 4 Thu',  'Pair Spotlight — GBPJPY'),
        ('Week 5 Tue',  'Trade Setup — GBPJPY Long (London Breakout + Retest)'),
        ('Week 5 Thu',  'Risk Rule #04 — Multi-Pair Diversification'),
        ('Week 6 Tue',  'Pair Spotlight — USDJPY'),
        ('Week 6 Thu',  'Trade Setup — USDJPY Short (Asian False Break)'),
    ],
    col_widths=(3.5, 12.5),
    header=('Slot', 'Content'),
)

heading2('5.2  Counter File')

code_block([
    '// data/edu-counter.json',
    '{ "index": 3 }',
    '',
    '// Counter advances by 1 after each post, wraps at 12.',
    '// Git-committed by GitHub Actions after every successful edu post.',
])

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 5. GITHUB ACTIONS WORKFLOW
# ════════════════════════════════════════════════════════════════

heading1('6. GitHub Actions Workflow')

heading2('6.1  Workflow File')

body('File: .github/workflows/insta-post.yml')

code_block([
    'Triggers:',
    '  schedule:',
    '    Mon 09:00 SGT  (01:00 UTC)  → weekly card',
    '    Tue 09:00 SGT  (01:00 UTC)  → edu post',
    '    Wed 09:00 SGT  (01:00 UTC)  → daily card',
    '    Thu 09:00 SGT  (01:00 UTC)  → edu post',
    '    Fri 09:00 SGT  (01:00 UTC)  → daily card',
    '    1st of month                → monthly chart',
    '  workflow_dispatch:            → manual trigger (POST_TYPE input)',
])

heading2('6.2  Workflow Steps')

table_2col(
    [
        ('1. Checkout',          'Checks out vera-level-fx repo — includes the latest vera-snapshot.json pushed by the local scraper.'),
        ('2. Set up Python',     'Installs Python 3.11 with pip cache keyed to instagram/requirements.txt.'),
        ('3. Install deps',      'pip install -r instagram/requirements.txt (matplotlib, Pillow, requests).'),
        ('4. Generate & post',   'Runs python instagram/run.py — determines post type, generates image, commits PNG, calls Meta API.'),
        ('5. Upload artifact',   'Saves the generated PNG as a GitHub Actions artifact (retained 30 days) for audit/review.'),
    ],
    col_widths=(4, 12),
    header=('Step', 'Action'),
)

heading2('6.3  Manual Trigger (workflow_dispatch)')

body(
    'Any post type can be triggered manually from GitHub → Actions → Instagram Auto-Post → Run workflow. '
    'The POST_TYPE input accepts:'
)

table_2col(
    [
        ('weekly',   'Weekly performance card (Monday format)'),
        ('daily',    'Daily live positions notepad card'),
        ('monthly',  'Monthly P&L bar chart'),
        ('trust',    'Win rate / trust card'),
        ('edu',      'Next educational post in the rotation sequence'),
    ],
    col_widths=(3, 13),
    header=('POST_TYPE value', 'Post generated'),
)

heading2('6.4  Required GitHub Secrets')

table_2col(
    [
        ('IG_USER_ID',        'Instagram Business account numeric ID (from Meta Developer dashboard)'),
        ('META_ACCESS_TOKEN', 'Long-lived Page Access Token — generated via Meta Graph API Explorer, valid 60 days, must be manually refreshed'),
        ('GITHUB_TOKEN',      'Auto-provided by GitHub Actions — used to commit the generated PNG back to the repo'),
    ],
    col_widths=(5, 11),
    header=('Secret', 'Purpose'),
)

note_box(
    '⏰  META_ACCESS_TOKEN expires every 60 days. Refresh it via the Meta Graph API Explorer '
    '(Extend Token button) and update the GitHub Secret before it expires — otherwise posting will fail silently.',
    'FFF3CD'
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 6. END-TO-END POSTING FLOW
# ════════════════════════════════════════════════════════════════

heading1('7. End-to-End Posting Flow (Step by Step)')

heading2('Phase 1 — Data Collection (Daily, 08:00 SGT, Windows Laptop)')

body('This phase runs locally regardless of what day it is.')
bullet('08:00 SGT — Windows Task Scheduler fires "node src/runner.js --push".')
bullet('Puppeteer launches stealth Chrome, navigates to Myfxbook Vera Level account page.')
bullet('Authenticates via Myfxbook login, scrapes #tradesListData for lifetime win rate, trade count, avg win/loss, pips, profit factor.')
bullet('Calls Myfxbook REST API (get-my-accounts, get-open-trades, get-open-orders, get-history, get-daily-gain) with authenticated session.')
bullet('Merges DOM-scraped lifetime stats with REST API data into a single snapshot object.')
bullet('Writes vera-level-fx/data/vera-snapshot.json and pushes to GitHub repo master branch.')
bullet('Telegram report sent to operator (balance, equity, open trades, daily gain, win rate).')
bullet('GitHub Actions in vera-level-fx repo now has access to the latest data.')

heading2('Phase 2 — Post Generation (Scheduled, 09:00 SGT, GitHub Actions)')

body('GitHub Actions picks up the fresh snapshot and generates the day\'s post.')
bullet('Cron fires at 01:00 UTC (09:00 SGT).')
bullet('Repo is checked out — data/vera-snapshot.json contains data from the 08:00 scrape.')
bullet('run.py loads vera-snapshot.json, determines post type from the day of week.')
bullet('Appropriate generator function is called with the live account data.')
bullet('Matplotlib renders the 1080×1080 PNG (no display server needed — Agg backend).')
bullet('PNG is saved to instagram/posts/YYYY-MM-DD-{type}.png.')
bullet('IC Markets referral QR code (camp=91936) is stamped onto the bottom-right corner via PIL.')
bullet('PNG is git-committed and pushed to the repo — creating a stable public URL on GitHub\'s CDN.')
bullet('20-second delay for CDN propagation.')
bullet('Caption is generated from the same live data.')
bullet('Meta Graph API is called: create media container → publish container.')
bullet('Post appears on @veralevel.fx Instagram feed.')
bullet('PNG uploaded as GitHub Actions artifact (30-day retention) for audit trail.')

heading2('Phase 3 — Educational Post (Tuesday / Thursday)')

body('Additional steps for educational posts:')
bullet('edu-counter.json is read to get the current rotation index (0–11).')
bullet('get_edu_content(index) returns the content type and data for that slot.')
bullet('The appropriate visual renderer is called (make_risk_post / make_pairs_post / make_setup_post).')
bullet('After publishing, counter is incremented (mod 12) and committed back to the repo.')
bullet('Next educational post will use index + 1, ensuring the 6-week sequence advances correctly.')

doc.add_paragraph()

heading2('Full Timing Summary (Weekday)')

table_2col(
    [
        ('08:00 SGT', 'Windows laptop: Myfxbook scraper runs, vera-snapshot.json updated and pushed to GitHub.'),
        ('08:00–09:00 SGT', 'GitHub syncs, CDN propagates the new data file across GitHub\'s edge network.'),
        ('09:00 SGT', 'GitHub Actions cron fires, runs the Instagram pipeline against the fresh snapshot.'),
        ('09:00–09:02 SGT', 'Image generated, committed, and CDN URL becomes available.'),
        ('09:02 SGT', 'Meta API: image container created + published. Post goes live on @veralevel.fx.'),
        ('09:02 SGT+', 'Artifact uploaded. Workflow completes. Instagram post is live.'),
    ],
    col_widths=(3.5, 12.5),
    header=('Time (SGT)', 'Event'),
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 7. SECURITY & CREDENTIALS
# ════════════════════════════════════════════════════════════════

heading1('8. Security & Credentials')

heading2('8.1  What Is Kept Secret')

table_2col(
    [
        ('IG_USER_ID',           'GitHub Actions Secret — Instagram Business account ID'),
        ('META_ACCESS_TOKEN',    'GitHub Actions Secret — Meta Graph API token (60-day expiry)'),
        ('MYFXBOOK_EMAIL',       'Local .env file in myfxbook-mcp (private repo) — Myfxbook login'),
        ('MYFXBOOK_PASSWORD',    'Local .env file in myfxbook-mcp (private repo) — Myfxbook login'),
        ('TELEGRAM_BOT_TOKEN',   'Local .env file in myfxbook-mcp (private repo) — Telegram bot'),
        ('TELEGRAM_CHAT_ID',     'Local .env file in myfxbook-mcp (private repo) — Telegram chat'),
    ],
    col_widths=(5, 11),
    header=('Credential', 'Storage location'),
)

heading2('8.2  What Is Public')

bullet('vera-level-fx GitHub repo — all Instagram code, image assets, vera-snapshot.json (account data only, no credentials).')
bullet('Generated PNG files committed to the repo — publicly visible via raw GitHub CDN URLs.')
bullet('edu-counter.json — the rotation index counter.')

heading2('8.3  Brand Name Protection')

body(
    'The system contains a hardcoded strip list to prevent certain brand names from appearing '
    'in any Instagram post content:'
)
code_block([
    '# generate_status.py',
    "_STRIP = ['tol langit', 'tol-langit', 'toklangit']",
    '',
    '# These strings must NEVER appear in any caption or image text.',
    '# The Myfxbook account (TOL LANGIT ETF GOLD) is a separate entity.',
    '# Vera Level FX is the only brand promoted on @veralevel.fx.',
])

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 8. TROUBLESHOOTING
# ════════════════════════════════════════════════════════════════

heading1('9. Troubleshooting')

heading2('9.1  Scraper / Data Issues')

table_2col(
    [
        ('All stats are zero in snapshot',
         'Cloudflare blocked the scrape. Disable VPN, check home IP is residential, run "node src/runner.js --dry" to test.'),
        ('"Just a moment" in accountName',
         'Cloudflare challenge page was returned. Same fix as above — wait and retry.'),
        ('Vera level API fetch failed',
         'Myfxbook session expired mid-run. Run runner.js again manually. Ensure MYFXBOOK_EMAIL and MYFXBOOK_PASSWORD are correct in .env.'),
        ('Git push failed after scrape',
         'Run: cd myfxbook-mcp && git pull --rebase && git push'),
        ('vera-snapshot.json not updating',
         'Check that VERA_FX_DIR path in runner.js points to the correct vera-level-fx directory on this machine.'),
    ],
    col_widths=(5, 11),
    header=('Symptom', 'Fix'),
)

heading2('9.2  Instagram Pipeline Issues')

table_2col(
    [
        ('GitHub Actions: no data',
         'The 08:00 SGT scrape may not have run yet, or the push failed. Check myfxbook-mcp git log. Manually trigger scrape if needed.'),
        ('Meta API 400 error',
         'Image URL not yet accessible on GitHub CDN. Increase sleep time in run.py (currently 20s) to 30s.'),
        ('Meta API 190 error (token invalid)',
         'META_ACCESS_TOKEN has expired. Go to Meta Graph API Explorer, extend the token, update GitHub Secret.'),
        ('Post generated but not published',
         'Check IG_USER_ID is correct — must be the numeric Business account ID, not the username.'),
        ('edu-counter stuck / wrong post type',
         'Manually edit data/edu-counter.json to the correct index (0–11) and commit.'),
        ('Monthly chart shows no bars',
         'dailyGain array is empty in vera-snapshot.json. The historyStart date in fetcher.js may need updating.'),
        ('Buffered image not being used',
         'Check the file is a .png (not .jpg or .PNG) and is in the correct subfolder. '
         'Run "ls instagram/buffer/daily/" to confirm. Also check git push succeeded so GitHub Actions can see the file.'),
        ('Wrong image used from buffer',
         'Buffer is FIFO — rename files with "01_", "02_" prefixes to control order. '
         'Or delete/archive unwanted files from the buffer folder and push again.'),
    ],
    col_widths=(5, 11),
    header=('Symptom', 'Fix'),
)

heading2('9.3  Manual Commands')

body('Test the scraper without sending Telegram or pushing:')
code_block(['cd "C:\\Users\\USER\\Downloads\\myfxbook-mcp"', 'node src/runner.js --dry'])

body('Force a specific Instagram post type:')
code_block([
    '# In GitHub Actions: Run workflow → set POST_TYPE to one of:',
    '# weekly | daily | monthly | trust | edu',
])

body('Generate a post locally for preview:')
code_block([
    'cd "C:\\Users\\USER\\Downloads\\04 - FX Trading\\vera-level-fx"',
    'python -c "',
    'import sys, json; sys.path.insert(0, \'instagram\')',
    'import matplotlib; matplotlib.use(\'Agg\')',
    'import matplotlib.pyplot as plt',
    'from generate_status import make_daily_card',
    'with open(\'data/vera-snapshot.json\') as f: data = json.load(f)',
    'fig = make_daily_card(data)',
    'fig.savefig(\'instagram/posts/preview.png\', dpi=100)',
    'plt.close(); print(\'Saved\')',
    '"',
])

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 9. MAINTENANCE CHECKLIST
# ════════════════════════════════════════════════════════════════

heading1('10. Maintenance Checklist')

heading2('Monthly')
bullet('Refresh META_ACCESS_TOKEN before 60-day expiry — update GitHub Secret.')
bullet('Review vera-snapshot.json data for any anomalies (all-zero fields, missing openTrades).')
bullet('Check GitHub Actions run history for failed jobs.')

heading2('Quarterly')
bullet('Review and refresh educational content in edu_content.py if market conditions have changed.')
bullet('Update example account sizes and stats in Risk Management rules if balance has grown significantly.')
bullet('Check IC Markets IB link is still valid: https://icmarkets.com/global/en/?camp=91936.')
bullet('Verify QR code (instagram/assets/qr-icmarkets.jpg) still scans to the correct referral URL.')

heading2('As Needed')
bullet('If Myfxbook changes their DOM selectors, update scraper.js CSS selectors.')
bullet('If Meta API changes Instagram publishing endpoints, update post.py.')
bullet('If background photo (bg-daily-notepad.jpg) needs updating, replace the asset and verify overlay positions still align in generate_status.py.')
bullet('Top up the content buffer (instagram/buffer/) with fresh pre-made images whenever you have design time — aim to keep at least 2–3 images queued per post type.')
bullet('Archive used buffer images from instagram/posts/*-buffered.png if you want to review what was posted.')

divider()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Vera Level FX  ·  @veralevel.fx  ·  Not financial advice')
set_run_font(run, size=9, color=MUTED, italic=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('IC Markets · ASIC Regulated  ·  All trades verified on Myfxbook  ·  Singapore')
set_run_font(run, size=9, color=MUTED, italic=True)


# ── Save ──────────────────────────────────────────────────────────
doc.save(OUT)
print(f'Saved: {OUT}')
